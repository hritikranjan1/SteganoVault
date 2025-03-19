from flask import Flask, request, send_file, jsonify, render_template
import os
import tempfile
import shutil
from PIL import Image
from stegano import lsb
from PyPDF2 import PdfReader, PdfWriter
from docx import Document
from pydub import AudioSegment
import cv2
import numpy as np

app = Flask(__name__)

def convert_to_png(input_path):
    """Convert image to PNG format."""
    try:
        img = Image.open(input_path)
        png_path = input_path.rsplit(".", 1)[0] + ".png"
        img.save(png_path, format="PNG")
        return png_path
    except Exception as e:
        print(f"Image Conversion Error: {e}")
        return None

def encode_image(input_path, message, password):
    """Embed a message in an image."""
    try:
        if not input_path.lower().endswith(".png"):
            input_path = convert_to_png(input_path)
        if not input_path:
            return None
        
        secret_message = f"{password}:{message}" if password else message
        encoded_img = lsb.hide(input_path, secret_message)
        output_path = input_path.replace(".", "_encoded.", 1)
        encoded_img.save(output_path)
        return output_path
    except Exception as e:
        print(f"Image Encoding Error: {e}")
        return None

def decode_image(input_path, password):
    """Extract a hidden message from an image."""
    try:
        extracted_message = lsb.reveal(input_path)
        if extracted_message:
            if ":" in extracted_message:
                stored_password, stored_message = extracted_message.split(":", 1)
                return stored_message if stored_password == password else "❌ Incorrect password!"
            return extracted_message
        return "❌ No hidden message found!"
    except Exception as e:
        print(f"Image Decoding Error: {e}")
        return "❌ Error decoding image!"

def encode_txt(input_path, message, password):
    """Embed a hidden message in a text file using zero-width characters."""
    try:
        with open(input_path, "r", encoding="utf-8") as file:
            original_content = file.read()
        
        secret_message = f"{password}:{message}" if password else message
        binary_data = ''.join(format(ord(c), '08b') for c in secret_message)
        encoded_message = ''.join("\u200B" if bit == "0" else "\u200D" for bit in binary_data)

        output_path = input_path.replace(".", "_encoded.", 1)
        with open(output_path, "w", encoding="utf-8") as file:
            file.write(original_content + encoded_message)
        
        return output_path
    except Exception as e:
        print(f"TXT Encoding Error: {e}")
        return None

def decode_txt(input_path, password):
    """Extract a hidden message from a text file using zero-width characters."""
    try:
        with open(input_path, "r", encoding="utf-8") as file:
            content = file.read()
        
        binary_data = ''.join("0" if char == "\u200B" else "1" for char in content if char in ["\u200B", "\u200D"])
        extracted_message = ''.join(chr(int(binary_data[i:i+8], 2)) for i in range(0, len(binary_data), 8))

        if extracted_message and ":" in extracted_message:
            stored_password, stored_message = extracted_message.split(":", 1)
            return stored_message if stored_password == password else "❌ Incorrect password!"
        
        return extracted_message if extracted_message else "❌ No hidden message found!"
    except Exception as e:
        print(f"TXT Decoding Error: {e}")
        return "❌ Error decoding TXT!"

def encode_docx(input_path, message, password):
    """Embed a hidden message in a DOCX file using hidden text formatting."""
    try:
        doc = Document(input_path)
        secret_message = f"{password}:{message}" if password else message
        paragraph = doc.add_paragraph()
        run = paragraph.add_run(secret_message)
        run.font.hidden = True
        
        output_path = input_path.replace(".", "_encoded.", 1)
        doc.save(output_path)
        return output_path
    except Exception as e:
        print(f"DOCX Encoding Error: {e}")
        return None

def decode_docx(input_path, password):
    """Extract a hidden message from a DOCX file."""
    try:
        doc = Document(input_path)
        extracted_message = ""
        for para in doc.paragraphs:
            for run in para.runs:
                if run.font.hidden:
                    extracted_message += run.text
        
        if extracted_message and ":" in extracted_message:
            stored_password, stored_message = extracted_message.split(":", 1)
            return stored_message if stored_password == password else "❌ Incorrect password!"
        
        return extracted_message if extracted_message else "❌ No hidden message found!"
    except Exception as e:
        print(f"DOCX Decoding Error: {e}")
        return "❌ Error decoding DOCX!"

def encode_pdf(input_path, message, password):
    """Embed a hidden message in a PDF file."""
    try:
        reader = PdfReader(input_path)
        writer = PdfWriter()
        for page in reader.pages:
            writer.add_page(page)
        metadata = {"/Message": f"{password}:{message}" if password else message}
        writer.add_metadata(metadata)
        
        output_path = input_path.replace(".", "_encoded.", 1)
        with open(output_path, "wb") as output_pdf:
            writer.write(output_pdf)
        return output_path
    except Exception as e:
        print(f"PDF Encoding Error: {e}")
        return None

def decode_pdf(input_path, password):
    """Extract a hidden message from a PDF file."""
    try:
        reader = PdfReader(input_path)
        metadata = reader.metadata
        extracted_message = metadata.get("/Message", "")
        
        if extracted_message and ":" in extracted_message:
            stored_password, stored_message = extracted_message.split(":", 1)
            return stored_message if stored_password == password else "❌ Incorrect password!"
        return extracted_message if extracted_message else "❌ No hidden message found!"
    except Exception as e:
        print(f"PDF Decoding Error: {e}")
        return "❌ Error decoding PDF!"

def convert_to_wav(input_path):
    """Convert audio file to WAV format."""
    try:
        audio = AudioSegment.from_file(input_path)
        wav_path = input_path.rsplit(".", 1)[0] + ".wav"
        audio.export(wav_path, format="wav")
        return wav_path
    except Exception as e:
        print(f"Audio Conversion Error: {e}")
        return None

def encode_audio(input_path, message, password):
    """Embed a hidden message in an audio file."""
    try:
        if not input_path.lower().endswith(".wav"):
            input_path = convert_to_wav(input_path)
        if not input_path:
            return None
        
        secret_message = f"{password}:{message}" if password else message
        binary_message = ''.join(format(ord(c), '08b') for c in secret_message)
        
        audio = AudioSegment.from_file(input_path)
        samples = audio.get_array_of_samples()
        
        if len(binary_message) > len(samples):
            return "❌ Message too large for the audio file!"
        
        for i in range(len(binary_message)):
            samples[i] = samples[i] & ~1 | int(binary_message[i])
        
        encoded_audio = audio._spawn(samples)
        output_path = input_path.replace(".", "_encoded.", 1)
        encoded_audio.export(output_path, format="wav")
        return output_path
    except Exception as e:
        print(f"Audio Encoding Error: {e}")
        return None

def decode_audio(input_path, password):
    """Extract a hidden message from an audio file."""
    try:
        if not input_path.lower().endswith(".wav"):
            input_path = convert_to_wav(input_path)
        if not input_path:
            return "❌ Error decoding audio!"

        audio = AudioSegment.from_file(input_path)
        samples = audio.get_array_of_samples()
        
        binary_message = ''.join(str(sample & 1) for sample in samples)
        extracted_message = ''
        for i in range(0, len(binary_message), 8):
            byte = binary_message[i:i+8]
            if not byte:
                break
            extracted_message += chr(int(byte, 2))
        
        # Remove trailing null characters
        extracted_message = extracted_message.rstrip('\x00')
        
        if extracted_message and ":" in extracted_message:
            stored_password, stored_message = extracted_message.split(":", 1)
            return stored_message if stored_password == password else "❌ Incorrect password!"
        
        return extracted_message if extracted_message else "❌ No hidden message found!"
    except Exception as e:
        print(f"Audio Decoding Error: {e}")
        return "❌ Error decoding audio!"

def encode_video(input_path, message, password):
    """Embed a hidden message in a video file."""
    try:
        secret_message = f"{password}:{message}" if password else message
        binary_message = ''.join(format(ord(c), '08b') for c in secret_message)
        
        # Open the video file
        cap = cv2.VideoCapture(input_path)
        frame_width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
        frame_height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
        fps = int(cap.get(cv2.CAP_PROP_FPS))
        fourcc = cv2.VideoWriter_fourcc(*'mp4v')  # Use 'mp4v' for MP4 format
        output_path = input_path.replace(".", "_encoded.", 1)
        out = cv2.VideoWriter(output_path, fourcc, fps, (frame_width, frame_height))
        
        frame_index = 0
        while cap.isOpened():
            ret, frame = cap.read()
            if not ret:
                break
            
            # Embed the message in the least significant bit of the blue channel
            if frame_index < len(binary_message):
                for i in range(frame.shape[0]):
                    for j in range(frame.shape[1]):
                        if frame_index < len(binary_message):
                            # Modify only the blue channel (0th index in OpenCV's BGR format)
                            frame[i, j, 0] = (frame[i, j, 0] & ~1) | int(binary_message[frame_index])
                            frame_index += 1
            
            out.write(frame)
        
        cap.release()
        out.release()
        return output_path
    except Exception as e:
        print(f"Video Encoding Error: {e}")
        return None

def decode_video(input_path, password):
    """Extract a hidden message from a video file."""
    try:
        cap = cv2.VideoCapture(input_path)
        binary_message = ''
        
        while cap.isOpened():
            ret, frame = cap.read()
            if not ret:
                break
            
            # Extract the message from the least significant bit of the blue channel
            for i in range(frame.shape[0]):
                for j in range(frame.shape[1]):
                    binary_message += str(frame[i, j, 0] & 1)
        
        cap.release()
        
        extracted_message = ''
        for i in range(0, len(binary_message), 8):
            byte = binary_message[i:i+8]
            if not byte:
                break
            extracted_message += chr(int(byte, 2))
        
        # Remove trailing null characters
        extracted_message = extracted_message.rstrip('\x00')
        
        if extracted_message and ":" in extracted_message:
            stored_password, stored_message = extracted_message.split(":", 1)
            return stored_message if stored_password == password else "❌ Incorrect password!"
        
        return extracted_message if extracted_message else "❌ No hidden message found!"
    except Exception as e:
        print(f"Video Decoding Error: {e}")
        return "❌ Error decoding video!"

@app.route("/")
def index():
    return render_template("index.html")

@app.route("/encode", methods=["POST"])
def encode():
    uploaded_file = request.files.get("file")
    message = request.form.get("message")
    password = request.form.get("password", "")

    if not uploaded_file or not message:
        return jsonify({"error": "File and message are required!"}), 400

    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, uploaded_file.filename)
    uploaded_file.save(file_path)

    ext = os.path.splitext(file_path)[1].lower()
    output_file = None
    
    try:
        if ext in [".png", ".jpg", ".jpeg"]:
            output_file = encode_image(file_path, message, password)
        elif ext == ".txt":
            output_file = encode_txt(file_path, message, password)
        elif ext == ".pdf":
            output_file = encode_pdf(file_path, message, password)
        elif ext == ".docx":
            output_file = encode_docx(file_path, message, password)
        elif ext in [".wav", ".mp3"]:
            output_file = encode_audio(file_path, message, password)
        elif ext in [".mp4", ".avi", ".mov"]:
            output_file = encode_video(file_path, message, password)
        else:
            return jsonify({"error": "Unsupported file format!"}), 400

        if output_file and os.path.exists(output_file):
            return send_file(output_file, as_attachment=True)
        return jsonify({"error": "Failed to encode file!"}), 500
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)

@app.route("/decode", methods=["POST"])
def decode():
    uploaded_file = request.files.get("file")
    password = request.form.get("password", "")

    if not uploaded_file:
        return jsonify({"error": "File is required!"}), 400

    temp_dir = tempfile.mkdtemp()
    file_path = os.path.join(temp_dir, uploaded_file.filename)
    uploaded_file.save(file_path)

    ext = os.path.splitext(file_path)[1].lower()
    decoded_message = None
    
    try:
        if ext in [".png", ".jpg", ".jpeg"]:
            decoded_message = decode_image(file_path, password)
        elif ext == ".txt":
            decoded_message = decode_txt(file_path, password)
        elif ext == ".pdf":
            decoded_message = decode_pdf(file_path, password)
        elif ext == ".docx":
            decoded_message = decode_docx(file_path, password)
        elif ext in [".wav", ".mp3"]:
            decoded_message = decode_audio(file_path, password)
        elif ext in [".mp4", ".avi", ".mov"]:
            decoded_message = decode_video(file_path, password)
        else:
            return jsonify({"error": "Unsupported file format!"}), 400

        return jsonify({"message": decoded_message})
    finally:
        shutil.rmtree(temp_dir, ignore_errors=True)

if __name__ == "__main__":
    app.run(debug=True)